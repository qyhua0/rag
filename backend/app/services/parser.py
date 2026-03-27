"""
文档解析服务 - 支持 PDF/Word/Excel/TXT/图片(OCR) 等格式
Package: top.modelx.rag
Author: hua
"""
import os
import io
import chardet
from pathlib import Path
from typing import List, Optional, Tuple
from loguru import logger

# LangChain document loaders
from langchain_core.documents import Document


class DocumentParser:
    """通用文档解析器"""

    SUPPORTED_TYPES = {
        "pdf": ["pdf"],
        "word": ["doc", "docx"],
        "excel": ["xls", "xlsx"],
        "text": ["txt", "md", "csv", "json", "xml", "html", "htm"],
        "image": ["jpg", "jpeg", "png", "gif", "bmp", "tiff", "webp"],
    }

    def __init__(self):
        self._ocr = None  # lazy load

    def _get_ocr(self):
        if self._ocr is None:
            try:
                from paddleocr import PaddleOCR
                self._ocr = PaddleOCR(use_angle_cls=True, lang="ch", show_log=False)
                logger.info("PaddleOCR initialized")
            except ImportError:
                logger.warning("PaddleOCR not available, image parsing disabled")
        return self._ocr

    def get_file_type(self, filename: str) -> str:
        ext = Path(filename).suffix.lower().lstrip(".")
        for ftype, exts in self.SUPPORTED_TYPES.items():
            if ext in exts:
                return ftype
        return "unknown"

    def parse(self, file_path: str, filename: str = None) -> Tuple[List[Document], dict]:
        """
        解析文件，返回 (documents, metadata)
        """
        filename = filename or os.path.basename(file_path)
        ext = Path(filename).suffix.lower().lstrip(".")
        file_type = self.get_file_type(filename)

        logger.info(f"Parsing file: {filename} (type={file_type})")

        try:
            if file_type == "pdf":
                docs = self._parse_pdf(file_path)
            elif file_type == "word":
                docs = self._parse_word(file_path)
            elif file_type == "excel":
                docs = self._parse_excel(file_path)
            elif file_type == "text":
                docs = self._parse_text(file_path)
            elif file_type == "image":
                docs = self._parse_image(file_path)
            else:
                raise ValueError(f"Unsupported file type: {ext}")

            total_chars = sum(len(d.page_content) for d in docs)
            meta = {
                "filename": filename,
                "file_type": file_type,
                "ext": ext,
                "page_count": len(docs),
                "total_chars": total_chars,
            }
            logger.info(f"Parsed {len(docs)} pages, {total_chars} chars from {filename}")
            return docs, meta

        except Exception as e:
            logger.error(f"Error parsing {filename}: {e}")
            raise

    def _parse_pdf(self, file_path: str) -> List[Document]:
        """解析 PDF 文件，文本页直接提取，图片页用 OCR"""
        from pypdf import PdfReader
        docs = []
        reader = PdfReader(file_path)

        for i, page in enumerate(reader.pages):
            text = page.extract_text() or ""
            text = text.strip()

            # 如果文本为空，尝试 OCR
            if not text:
                try:
                    import fitz  # PyMuPDF fallback
                    pdf_doc = fitz.open(file_path)
                    pdf_page = pdf_doc[i]
                    pix = pdf_page.get_pixmap(dpi=200)
                    img_bytes = pix.tobytes("png")
                    text = self._ocr_bytes(img_bytes)
                    pdf_doc.close()
                except Exception:
                    pass

            if text:
                docs.append(Document(
                    page_content=text,
                    metadata={"source": file_path, "page": i + 1}
                ))

        return docs if docs else [Document(page_content="", metadata={"source": file_path})]

    def _parse_word(self, file_path: str) -> List[Document]:
        """解析 Word 文档 (.doc/.docx)"""
        ext = Path(file_path).suffix.lower()
        if ext == ".docx":
            from docx import Document as DocxDocument
            doc = DocxDocument(file_path)
            paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            # 合并表格内容
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        paragraphs.append(row_text)
            content = "\n".join(paragraphs)
        else:
            # .doc 格式使用 antiword 或 python-docx2txt
            try:
                import docx2txt
                content = docx2txt.process(file_path)
            except Exception:
                content = f"[无法解析 .doc 文件: {file_path}]"

        return [Document(page_content=content, metadata={"source": file_path})]

    def _parse_excel(self, file_path: str) -> List[Document]:
        """解析 Excel 文件 (.xls/.xlsx)"""
        import openpyxl
        docs = []
        ext = Path(file_path).suffix.lower()

        if ext == ".xlsx":
            wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                rows = []
                for row in ws.iter_rows(values_only=True):
                    row_data = [str(cell) if cell is not None else "" for cell in row]
                    if any(cell.strip() for cell in row_data):
                        rows.append(" | ".join(row_data))
                content = f"[Sheet: {sheet_name}]\n" + "\n".join(rows)
                docs.append(Document(
                    page_content=content,
                    metadata={"source": file_path, "sheet": sheet_name}
                ))
            wb.close()
        else:
            # .xls 格式
            import xlrd
            wb = xlrd.open_workbook(file_path)
            for sheet in wb.sheets():
                rows = []
                for rx in range(sheet.nrows):
                    row_data = [str(sheet.cell_value(rx, cx)) for cx in range(sheet.ncols)]
                    rows.append(" | ".join(row_data))
                content = f"[Sheet: {sheet.name}]\n" + "\n".join(rows)
                docs.append(Document(
                    page_content=content,
                    metadata={"source": file_path, "sheet": sheet.name}
                ))

        return docs

    def _parse_text(self, file_path: str) -> List[Document]:
        """解析纯文本文件，自动检测编码"""
        with open(file_path, "rb") as f:
            raw = f.read()

        detected = chardet.detect(raw)
        encoding = detected.get("encoding") or "utf-8"

        try:
            content = raw.decode(encoding, errors="replace")
        except Exception:
            content = raw.decode("utf-8", errors="replace")

        return [Document(page_content=content, metadata={"source": file_path})]

    def _parse_image(self, file_path: str) -> List[Document]:
        """使用 PaddleOCR 解析图片"""
        ocr = self._get_ocr()
        if ocr is None:
            return [Document(page_content="[OCR不可用，无法解析图片]", metadata={"source": file_path})]

        try:
            result = ocr.ocr(file_path, cls=True)
            text_lines = []
            if result:
                for line in result:
                    if line:
                        for item in line:
                            if item and len(item) >= 2:
                                text_lines.append(item[1][0])
            content = "\n".join(text_lines) if text_lines else "[图片中未识别到文字]"
        except Exception as e:
            logger.error(f"OCR error: {e}")
            content = f"[OCR解析失败: {str(e)}]"

        return [Document(page_content=content, metadata={"source": file_path, "ocr": True})]

    def _ocr_bytes(self, img_bytes: bytes) -> str:
        """对图片字节数据进行 OCR"""
        ocr = self._get_ocr()
        if ocr is None:
            return ""
        try:
            import numpy as np
            from PIL import Image
            img = Image.open(io.BytesIO(img_bytes))
            img_array = np.array(img)
            result = ocr.ocr(img_array, cls=True)
            lines = []
            if result:
                for line in result:
                    if line:
                        for item in line:
                            if item and len(item) >= 2:
                                lines.append(item[1][0])
            return "\n".join(lines)
        except Exception as e:
            logger.error(f"OCR bytes error: {e}")
            return ""


# 全局解析器单例
parser = DocumentParser()
